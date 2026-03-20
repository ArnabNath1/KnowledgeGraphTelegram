"""
Document Parser - handles PDF, text, images, markdown
"""
import io
import os
import tempfile
from pathlib import Path
from loguru import logger


class DocumentParser:
    """Parse various document formats into plain text"""

    async def parse(self, file_bytes: bytes, filename: str) -> str:
        """Route parsing based on file extension"""
        ext = Path(filename).suffix.lower()
        logger.info(f"Parsing document: {filename} (type: {ext})")

        try:
            if ext == ".pdf":
                return await self.parse_pdf(file_bytes)
            elif ext in (".txt", ".md", ".tex"):
                return file_bytes.decode("utf-8", errors="ignore")
            elif ext in (".docx", ".doc"):
                return await self.parse_docx(file_bytes)
            elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"):
                return await self.parse_image(file_bytes)
            else:
                # Try as plain text
                return file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Parse error for {filename}: {e}")
            raise

    async def parse_pdf(self, file_bytes: bytes) -> str:
        """Parse PDF using pdfplumber (more reliable) with PyMuPDF fallback"""
        text = ""
        # Try pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(f"[Page {i+1}]\n{page_text}")
                text = "\n\n".join(pages)
                logger.info(f"pdfplumber extracted {len(text)} chars from {len(pdf.pages)} pages")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyMuPDF")

        # Fallback to PyMuPDF
        if not text.strip():
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                pages = []
                for i, page in enumerate(doc):
                    page_text = page.get_text()
                    if page_text.strip():
                        pages.append(f"[Page {i+1}]\n{page_text}")
                text = "\n\n".join(pages)
                logger.info(f"PyMuPDF extracted {len(text)} chars")
            except Exception as e:
                logger.error(f"PyMuPDF also failed: {e}")
                raise Exception("Could not extract text from PDF")

        if not text.strip():
            raise Exception("PDF appears to be image-only. Try OCR or copy-paste text.")
        return text

    async def parse_docx(self, file_bytes: bytes) -> str:
        """Parse .docx files"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Include table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            raise

    async def parse_image(self, file_bytes: bytes) -> str:
        """OCR image using pytesseract"""
        try:
            from PIL import Image
            import pytesseract
            img = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(img, lang="eng")
            if text.strip():
                return text
            raise Exception("No text found in image")
        except Exception as e:
            logger.error(f"OCR error: {e}")
            raise Exception(f"Could not extract text from image: {e}")

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        import re
        # Remove excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        # Remove page headers/footers patterns
        text = re.sub(r"^\d+\s*$", "", text, flags=re.MULTILINE)
        return text.strip()
